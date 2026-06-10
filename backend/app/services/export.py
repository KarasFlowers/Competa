"""Report export service — convert report JSON to Markdown or Word (.docx)."""

from __future__ import annotations

import io
from typing import Any


def _has_curation_decisions(sources: list[dict[str, Any]]) -> bool:
    return any(
        source.get("included_in_analysis")
        or source.get("curation_reason")
        or source.get("curated_excerpt")
        or source.get("curation_tags")
        for source in sources
    )


def _split_sources(
    sources: list[dict[str, Any]],
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    if not _has_curation_decisions(sources):
        return sources, []
    included = [source for source in sources if source.get("included_in_analysis")]
    excluded = [source for source in sources if not source.get("included_in_analysis")]
    return included, excluded


def _format_percentage(value: Any) -> str:
    try:
        return f"{round(float(value) * 100)}%"
    except Exception:
        return ""


def _curation_reason_label(reason: str) -> str:
    labels = {
        "selected": "Included in analysis",
        "duplicate_url": "Removed as duplicate URL",
        "duplicate_content": "Removed as duplicate content",
        "low_reliability": "Removed for low reliability",
        "domain_cap": "Removed due to domain diversity cap",
        "max_source_cap": "Removed due to source cap",
    }
    return labels.get(reason, reason.replace("_", " ").strip() or "Not labeled")


def _source_detail_lines(source: dict[str, Any], *, has_curation: bool) -> list[str]:
    lines: list[str] = []
    reliability = _format_percentage(source.get("reliability_score"))
    if reliability:
        lines.append(f"Reliability: {reliability}")
    if source.get("url"):
        lines.append(f"URL: {source['url']}")
    if has_curation:
        lines.append(f"Curation: {_curation_reason_label(str(source.get('curation_reason', '')))}")
    excerpt = source.get("curated_excerpt") or source.get("content_snippet") or ""
    if excerpt:
        lines.append(f"Excerpt: {str(excerpt)[:240]}")
    tags = [str(tag) for tag in source.get("curation_tags", []) if tag]
    if tags and has_curation:
        lines.append(f"Tags: {', '.join(tags[:5])}")
    return lines


def report_to_markdown(report: dict[str, Any], sources: list[dict[str, Any]] | None = None) -> str:
    """Convert a report dict to a Markdown string."""
    lines: list[str] = []

    # Title
    title = report.get("title", "Competitive Analysis Report")
    lines.append(f"# {title}")
    lines.append("")

    # Executive summary
    summary = report.get("executive_summary", "")
    if summary:
        lines.append("## Executive Summary")
        lines.append("")
        lines.append(summary)
        lines.append("")

    # Sections (recursive)
    def _render_sections(sections: list[dict], depth: int = 2) -> None:
        for section in sections:
            heading = "#" * (depth + 1) + " " + section.get("title", "Untitled")
            lines.append(heading)
            lines.append("")

            content = section.get("content", "")
            if content:
                lines.append(content)
                lines.append("")

            # Claims
            for claim in section.get("claims", []):
                claim_text = claim.get("content", "")
                evidence_ids = claim.get("evidence_ids", [])
                severity = claim.get("severity", "")
                prefix = "- "
                if severity:
                    prefix = f"- **[{severity.upper()}]** "
                line = f"{prefix}{claim_text}"
                if evidence_ids:
                    line += f" *(evidence: {', '.join(str(e) for e in evidence_ids)})*"
                lines.append(line)
            if section.get("claims"):
                lines.append("")

            # Subsections
            _render_sections(section.get("subsections", []), depth + 1)

    _render_sections(report.get("sections", []))

    # Source references
    if sources:
        has_curation = _has_curation_decisions(sources)
        included_sources, excluded_sources = _split_sources(sources)

        lines.append("## Sources Used in Analysis" if has_curation else "## Sources")
        lines.append("")
        for i, src in enumerate(included_sources, 1):
            src_type = src.get("type", "url")
            title = src.get("title", "Untitled")
            lines.append(f"[{i}] **{title}** ({src_type})")
            for detail_line in _source_detail_lines(src, has_curation=has_curation):
                lines.append(f"    {detail_line}")
            lines.append("")

        if excluded_sources:
            lines.append("## Excluded or Deprioritized Sources")
            lines.append("")
            for i, src in enumerate(excluded_sources, len(included_sources) + 1):
                src_type = src.get("type", "url")
                title = src.get("title", "Untitled")
                lines.append(f"[{i}] **{title}** ({src_type})")
                for detail_line in _source_detail_lines(src, has_curation=True):
                    lines.append(f"    {detail_line}")
                lines.append("")

    return "\n".join(lines)


def report_to_docx(report: dict[str, Any], sources: list[dict[str, Any]] | None = None) -> bytes:
    """Convert a report dict to a Word (.docx) file and return bytes."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        raise RuntimeError("python-docx is not installed. Install with: pip install python-docx")

    doc = Document()

    # Title
    title = report.get("title", "Competitive Analysis Report")
    doc.add_heading(title, level=0)

    # Executive summary
    summary = report.get("executive_summary", "")
    if summary:
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(summary)

    # Sections (recursive)
    def _add_sections(sections: list[dict], level: int = 1) -> None:
        for section in sections:
            doc.add_heading(section.get("title", "Untitled"), level=min(level, 9))
            content = section.get("content", "")
            if content:
                doc.add_paragraph(content)

            for claim in section.get("claims", []):
                claim_text = claim.get("content", "")
                severity = claim.get("severity", "")
                prefix = f"[{severity.upper()}] " if severity else ""
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(f"{prefix}{claim_text}")
                evidence_ids = claim.get("evidence_ids", [])
                if evidence_ids:
                    ev_run = p.add_run(f" (evidence: {', '.join(str(e) for e in evidence_ids)})")
                    ev_run.font.color.rgb = RGBColor(128, 128, 128)
                    ev_run.font.size = Pt(9)

            _add_sections(section.get("subsections", []), level + 1)

    _add_sections(report.get("sections", []))

    # Sources
    if sources:
        has_curation = _has_curation_decisions(sources)
        included_sources, excluded_sources = _split_sources(sources)

        doc.add_heading("Sources Used in Analysis" if has_curation else "Sources", level=1)
        for i, src in enumerate(included_sources, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"[{i}] {src.get('title', 'Untitled')} ({src.get('type', 'url')})")
            run.bold = True
            for detail_line in _source_detail_lines(src, has_curation=has_curation):
                doc.add_paragraph(detail_line)

        if excluded_sources:
            doc.add_heading("Excluded or Deprioritized Sources", level=1)
            for i, src in enumerate(excluded_sources, len(included_sources) + 1):
                p = doc.add_paragraph()
                run = p.add_run(f"[{i}] {src.get('title', 'Untitled')} ({src.get('type', 'url')})")
                run.bold = True
                for detail_line in _source_detail_lines(src, has_curation=True):
                    doc.add_paragraph(detail_line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
